#CREAR 25 ARCHIVOS WORD#
import random, os
import shutil
from docx import Document
from docx.shared import Cm

nombres = (
    'Final_Poo_tarea_',
    'Labs_POO_tarea_',
    'Telecomunicaciones1_tarea_',
    'Telecomunicaiones2_tarea_',
    'lab_Tele3_tarea_',
    'lab_tele4_tarea_',
    'Programacion1_tarea_',
    'IngenierriaaSocial_tarea_',
    'SistemasOperativaos_tarea_',
    'SDS_tarea_',
    'segBasesDatos_tarea_',
    'BasesDatos_tarea_',
    'SEGA_tarea_',
    'LabSEGA_tarea_',
    'Coco_tarea_',
    'Criptografia_tarea_',
    'DiseñoArquitecturas_tarea_',
    'NTSI_tarea_',
    'SegCompNube_tarea_'
)

def elegir_materia():
    materia = random.choice(nombres)
    materia.encode()
    return materia

def main():
    for iteracion in range(1, 26):
        # Creación del documento
        document = Document()

        # Añadimos un titulo al documento, a nivel 0
        titulo_documento = "Documento"+str(iteracion)+" creado con Python"
        document.add_heading(titulo_documento, 0)

        # Añadimos un párrafo
        p = document.add_paragraph('El contenido de los párrafos se añadir en varias líneas. ')
        p.add_run('Pudiéndose configurar que el texto tenga formato tipo ')
        p.add_run('negrita').bold = True
        p.add_run(' o ')
        p.add_run('itálica.').italic = True

        # Para indicar subtitulo se indica el nivel 1
        document.add_heading('Subtitulo', level=1)

        document.add_paragraph('Ahora se puede crear una enumeración')
        document.add_paragraph('Uno', style='List Number')
        document.add_paragraph('Dos', style='List Number')
        document.add_paragraph('Tres', style='List Number')

        document.add_paragraph('O viñetas')
        document.add_paragraph('Manzana', style='List Bullet') 
        document.add_paragraph('Pera', style='List Bullet')
        document.add_paragraph('Naranja', style='List Bullet')
        document.add_paragraph('Piña', style='List Bullet') 
        document.add_paragraph('Melón', style='List Bullet')
        document.add_paragraph('Sandía', style='List Bullet')

        # Tablas
        document.add_heading('Tablas', level=1)

        data = (('Manzana', 12), ('Pera', 5), ('Naranja', 12), ('Piña', 89), ('Melón', 65), ('Sandía', 7))

        table = document.add_table(rows=1, cols=2)

        table.rows[0].cells[0].text = 'Fruta'
        table.rows[0].cells[1].text = 'Cantidad'

        for prod, numbr in data:
            row_cells = table.add_row().cells
            row_cells[0].text = prod
            row_cells[1].text = str(numbr)


        p = document.add_paragraph()
        r2 = p.add_run()

        document.add_paragraph('Credit Card')
        document.add_paragraph('CVV')

        archivos = os.listdir('./imgs')
        imagen_seleccionado = random.choice(archivos)

        r2.add_picture(f'./imgs/{imagen_seleccionado}')
        r2.add_picture(f'./imgs/{imagen_seleccionado}')


        nombre_documento = elegir_materia()
        nombre_documento = nombre_documento + str(iteracion)+".docx"

        document.save(nombre_documento)


if __name__ == '__main__':
	main()
print("Hecho!!!")